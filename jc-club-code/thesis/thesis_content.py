# -*- coding: utf-8 -*-
"""
JC-Club毕业设计论文生成脚本
基于DDD微服务架构的程序员社区平台的设计与实现
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ======== 页面设置 ========
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = Pt(22)

def add_paragraph(text, font_name='宋体', font_size=Pt(12), bold=False, alignment=None, space_after=Pt(0), space_before=Pt(0), first_line_indent=Cm(0.74)):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    return p

def add_title_center(text, font_name='宋体', font_size=Pt(16), bold=True):
    """添加居中标题"""
    return add_paragraph(text, font_name=font_name, font_size=font_size, bold=bold,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

def add_body(text):
    """添加正文段落"""
    return add_paragraph(text, font_name='宋体', font_size=Pt(12), bold=False,
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74))

def add_chapter_title(text):
    """添加章标题"""
    return add_paragraph(text, font_name='黑体', font_size=Pt(16), bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12),
                        first_line_indent=None)

def add_section_title(text):
    """添加节标题"""
    return add_paragraph(text, font_name='黑体', font_size=Pt(14), bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6), space_after=Pt(6),
                        first_line_indent=None)

def add_subsection_title(text):
    """添加小节标题"""
    return add_paragraph(text, font_name='黑体', font_size=Pt(12), bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6), space_after=Pt(6),
                        first_line_indent=None)

def add_code_block(text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    return p

def add_table_header(table, headers, widths=None):
    """设置表格头"""
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.bold = True

def add_table_row(table, row_idx, values):
    """填充表格行"""
    for i, val in enumerate(values):
        cell = table.rows[row_idx].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)

# ============================================================
# 封面页
# ============================================================
# 空行
for _ in range(2):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(32)

# 单位学号行
add_paragraph("单位 计科2104班", font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None)
add_paragraph("学号 21203436", font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None)

# 空行
for _ in range(1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(32)

add_paragraph("江西农业大学南昌商学院本科毕业论文", font_name='宋体', font_size=Pt(16), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("（物联网工程专业）", font_name='宋体', font_size=Pt(16), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

for _ in range(1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(32)

add_paragraph("基于DDD微服务架构的程序员社区平台", font_name='宋体', font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("jc-club的设计与实现", font_name='宋体', font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(32)

add_paragraph("姓 名 张铭宇", font_name='黑体', font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("专 业 物联网工程", font_name='黑体', font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("指导教师 徐颖慧", font_name='黑体', font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

for _ in range(2):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(32)

add_paragraph("江西农业大学南昌商学院", font_name='宋体', font_size=Pt(16), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("二〇二六年五月", font_name='宋体', font_size=Pt(16),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

# 分页
doc.add_page_break()

# ============================================================
# 论文独创性声明
# ============================================================
add_title_center("论文独创性声明", font_name='黑体', font_size=Pt(16))

doc.add_paragraph()
add_body("本人声明，所呈交的学位论文系在导师指导下独立完成的研究成果。文中合法应用他人的成果，均已做出明确标注或得到许可。论文内容未包含法律意义上已属于他人的任何形式的研究成果，也不包含本人已用于其他学位申请的论文或成果。")
doc.add_paragraph()
add_body("本文如违反上述声明，愿意承担以下责任和后果：")
add_body("1、交回学校授予的学士学位；")
add_body("2、学校可在相关媒体上对本人的行为进行通报；")
add_body("3、本文按照学校规定的方式，对因不当取得学位给学校造成的名誉损害，进行公开道歉；")
add_body("4、本人负责因论文成果不实产生的法律纠纷。")
doc.add_paragraph()
doc.add_paragraph()
add_paragraph("论文作者签名：  日期： 2026年 5 月 15 日", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None)

doc.add_page_break()

# ============================================================
# 中文摘要
# ============================================================
add_paragraph("江西农业大学南昌商学院", font_name='宋体', font_size=Pt(10),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
add_paragraph("摘要", font_name='黑体', font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(22)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

abstract_cn = (
    "本论文基于领域驱动设计（DDD）与Spring Cloud Alibaba微服务架构，采用前后端分离模式，"
    "设计并实现了一款面向程序员群体的综合性社区平台——jc-club。后端基于Spring Boot 3.x框架，"
    "运用Nacos服务注册与发现、Gateway网关路由、OpenFeign远程调用构建微服务体系，"
    "整合MySQL数据库、Redis缓存、Elasticsearch全文检索和RocketMQ消息队列等中间件技术，"
    "将系统拆分为题目管理、用户认证、在线练习、模拟面试、社区圈子和对象存储六个独立微服务模块。"
    "前端采用React 18与TypeScript技术栈，基于Vite构建工具进行工程化开发，"
    "实现了题库浏览与检索、专项练习与智能组卷、AI模拟面试、简历分析、"
    "动态发布与评论互动等核心功能。平台围绕程序员"学-练-测-面"的学习闭环进行设计，"
    "为开发者提供了一站式的技能提升与求职准备解决方案。测试结果表明，各微服务模块协同运行稳定，"
    "功能完整可用，系统在高并发场景下表现出良好的性能与可靠性，"
    "具有较高的实用价值和良好的扩展性。"
)
add_body(abstract_cn)
doc.add_paragraph()

add_paragraph("【关键词】在线刷题；模拟面试；社区互动；智能组卷", font_name='黑体', font_size=Pt(12),
              bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)

doc.add_paragraph()
add_paragraph("I", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

doc.add_page_break()

# ============================================================
# 英文摘要
# ============================================================
add_paragraph("江西农业大学南昌商学院", font_name='宋体', font_size=Pt(10),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
add_paragraph("Abstract", font_name='Times New Roman', font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(22)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

abstract_en = (
    "This paper presents a comprehensive programmer community platform, jc-club, "
    "designed and implemented based on Domain-Driven Design (DDD) and the Spring Cloud Alibaba "
    "microservice architecture with a front-end and back-end separation model. The backend, "
    "built on the Spring Boot 3.x framework, leverages Nacos for service registration and discovery, "
    "Gateway for unified routing, and OpenFeign for remote procedure calls to construct the microservice "
    "ecosystem. The architecture integrates MySQL for persistent storage, Redis for caching, "
    "Elasticsearch for full-text search, and RocketMQ for asynchronous messaging, decomposing the system "
    "into six independent microservices: subject management, user authentication, online practice, "
    "mock interview, community circle, and object storage. The frontend employs React 18 with TypeScript, "
    "utilizing Vite as the build tool for engineering development. The platform implements core features "
    "including question bank browsing and searching, targeted practice with intelligent test paper generation, "
    "AI-powered mock interviews, resume analysis, and community interaction with post publishing and commenting. "
    "Designed around the \"learn-practice-assess-interview\"闭环 for programmers, the platform provides "
    "a one-stop solution for skill enhancement and job preparation. Testing results demonstrate that "
    "all microservice modules operate stably with complete functionality, and the system exhibits "
    "strong performance and reliability under high concurrency scenarios, offering significant practical "
    "value and favorable extensibility."
)
p = add_body(abstract_en)
# 设置英文摘要字体
for run in p.runs:
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

add_paragraph("【Keywords】online practice；mock interview；community interaction；intelligent test paper generation",
              font_name='Times New Roman', font_size=Pt(12), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)

doc.add_paragraph()
add_paragraph("II", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

doc.add_page_break()

# ============================================================
# 目录页（简化版）
# ============================================================
add_paragraph("江西农业大学南昌商学院", font_name='宋体', font_size=Pt(10),
              alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
add_paragraph("目录", font_name='黑体', font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
doc.add_paragraph()

toc_text = """摘要 ....................................... I
Abstract .................................. II
1 绪论 ..................................... 1
   1.1 软件研究背景和意义....................................................................... 1
   1.2 软件的国内外的研究现状及发展趋势........................................... 2
      1.2.1 国内研究现状........................................................................ 2
      1.2.2 国外研究现状........................................................................ 3
   1.3 软件设计内容及目标....................................................................... 4
2 系统分析 ................................. 5
   2.1 需求分析........................................................................................... 5
   2.2 业务流程分析................................................................................... 6
   2.3 功能模块分析................................................................................... 7
   2.4 技术分析........................................................................................... 10
3 系统设计 ................................. 12
   3.1 软件总体设计................................................................................... 12
   3.2 系统详细设计................................................................................... 13
      3.2.1 普通用户功能模块设计.......................................................... 13
      3.2.2 管理员功能模块设计.............................................................. 16
   3.3 数据库的设计................................................................................... 18
      3.3.1 概要设计.................................................................................. 19
      3.3.2 逻辑设计.................................................................................. 24
4 系统实现 ................................. 30
   4.1 前端功能实现................................................................................... 30
      4.1.1 用户登录注册模块.................................................................. 30
      4.1.2 题库浏览与检索模块.............................................................. 32
      4.1.3 专项练习模块.......................................................................... 34
      4.1.4 模拟面试模块.......................................................................... 36
      4.1.5 社区圈子模块.......................................................................... 37
   4.2 后端功能实现................................................................................... 38
      4.2.1 题目管理模块.......................................................................... 38
      4.2.2 用户认证与权限管理模块...................................................... 42
      4.2.3 练习管理模块.......................................................................... 44
      4.2.4 面试管理模块.......................................................................... 45
      4.2.5 圈子管理模块.......................................................................... 46
      4.2.6 文件存储与网关模块.............................................................. 47
5 系统测试 ................................. 49
   5.1 系统测试概述................................................................................... 49
   5.2 系统测试过程................................................................................... 50
      5.2.1 测试用例.................................................................................. 50
   5.3 测试结果........................................................................................... 56
结论 ...................................... 57
参考文献 .................................. 58
附录 ...................................... 60
致谢 ...................................... 76"""

for line in toc_text.strip().split('\n'):
    p = add_paragraph(line, font_name='宋体', font_size=Pt(12), first_line_indent=None)
    p.paragraph_format.line_spacing = Pt(22)

doc.add_page_break()

print("封面、摘要、目录生成完成...")
doc.save(r'D:\code\jc-club\jc-club-code\thesis\jc-club-thesis-part1.docx')
print("Part1 已保存")
