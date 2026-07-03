# -*- coding: utf-8 -*-
"""
Generate the complete jc-club thesis docx from content files.
Reads Chinese content from .txt files to avoid encoding issues.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

CONTENT_DIR = os.path.join(os.path.dirname(__file__), 'content')

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'  # Song
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = Pt(22)

def read_content(filename):
    """Read content from a text file with UTF-8 encoding."""
    filepath = os.path.join(CONTENT_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read().strip()
    return ''

def add_para(text, font_east='宋体', font_west='Times New Roman',
             size=Pt(12), bold=False, align=None, first_indent=None,
             space_before=Pt(0), space_after=Pt(0), line_spacing=Pt(22)):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if first_indent is not None:
        p.paragraph_format.first_line_indent = first_indent
    run = p.add_run(text)
    run.font.name = font_west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_east)
    run.font.size = size
    run.bold = bold
    return p

def add_cover_text(text, size=Pt(14), font_east='宋体', bold=False):
    """Add cover page text."""
    return add_para(text, font_east=font_east, size=size, bold=bold,
                    align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

def add_body(text):
    """Add body paragraph with first-line indent."""
    return add_para(text, first_indent=Cm(0.74), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

def add_chapter_title(text):
    """Add chapter title."""
    return add_para(text, font_east='黑体', size=Pt(16), bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12),
                    space_after=Pt(12), first_indent=None)

def add_section_title(text):
    """Add section title."""
    return add_para(text, font_east='黑体', size=Pt(14), bold=True,
                    space_before=Pt(6), space_after=Pt(6), first_indent=None)

def add_subsection_title(text):
    """Add subsection title."""
    return add_para(text, font_east='黑体', size=Pt(12), bold=True,
                    space_before=Pt(6), space_after=Pt(6), first_indent=None)

def add_code(text):
    """Add code block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.first_line_indent = Cm(0)
    for line in text.strip().split('\n'):
        if line.strip():
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(8)
    return p

def add_table_with_data(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.bold = True
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_empty_line():
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ================================================================
# COVER PAGE
# ================================================================
add_empty_line()
add_empty_line()
add_para('', line_spacing=Pt(32))
add_para('单位 计科2104班', size=Pt(12),
         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('学号 21203436', size=Pt(12),
         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('', line_spacing=Pt(32))
add_para('江西农业大学南昌商学院本科毕业论文',
         font_east='宋体', size=Pt(16), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('（物联网工程专业）',
         font_east='宋体', size=Pt(16), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('', line_spacing=Pt(32))
add_cover_text('基于DDD微服务架构的程序员社区平台',
               size=Pt(22), bold=True)
add_cover_text('jc-club的设计与实现', size=Pt(22), bold=True)
add_para('', line_spacing=Pt(32))
add_para('', line_spacing=Pt(32))
add_para('', line_spacing=Pt(32))

add_cover_text('姓 名 张铭宇', font_east='黑体', size=Pt(14))
add_cover_text('专 业 物联网工程', font_east='黑体', size=Pt(14))
add_cover_text('指导教师 徐颖慧', font_east='黑体', size=Pt(14))

add_para('', line_spacing=Pt(32))
add_para('', line_spacing=Pt(32))

add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(16), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('二〇二六年五月',
         font_east='宋体', size=Pt(16),
         align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ================================================================
# DECLARATION PAGE
# ================================================================
add_chapter_title('论文独创性声明')
add_empty_line()
add_body('本人声明，所呈交的学位论文系在导师指导下独立完成的研究成果。文中合法应用他人的成果，均已做出明确标注或得到许可。论文内容未包含法律意义上已属于他人的任何形式的研究成果，也不包含本人已用于其他学位申请的论文或成果。')
add_empty_line()
add_body('本文如违反上述声明，愿意承担以下责任和后果：')
add_body('1、交回学校授予的学士学位；')
add_body('2、学校可在相关媒体上对本人的行为进行通报；')
add_body('3、本文按照学校规定的方式，对因不当取得学位给学校造成的名誉损害，进行公开道歉；')
add_body('4、本人负责因论文成果不实产生的法律纠纷。')
add_empty_line()
add_empty_line()
add_para('论文作者签名：  日期： 2026年 5 月 15 日',
         align=WD_ALIGN_PARAGRAPH.RIGHT)

page_break()

# ================================================================
# ABSTRACT (Chinese)
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('摘要', font_east='黑体', size=Pt(22), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_line()

abstract_text = read_content('abstract_cn.txt')
if abstract_text:
    add_body(abstract_text)

add_empty_line()
add_para('【关键词】在线刷题；模拟面试；社区互动；智能组卷',
         font_east='黑体', size=Pt(12), bold=True)
add_empty_line()
add_para('I', align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ================================================================
# ABSTRACT (English)
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('Abstract', font_west='Times New Roman', size=Pt(22), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_line()

abstract_en = read_content('abstract_en.txt')
if abstract_en:
    add_body(abstract_en)

add_empty_line()
add_para('【Keywords】online practice；mock interview；community interaction；intelligent test paper generation',
         font_west='Times New Roman', size=Pt(12), bold=True)
add_empty_line()
add_para('II', align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ================================================================
# TABLE OF CONTENTS
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_para('目录', font_east='黑体', size=Pt(22), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_line()

toc_lines = [
    '摘要 ....................................... I',
    'Abstract .................................. II',
    '1 绪论 ..................................... 1',
    '   1.1 软件研究背景和意义............................................... 1',
    '   1.2 软件的国内外的研究现状及发展趋势................... 2',
    '      1.2.1 国内研究现状................................................ 2',
    '      1.2.2 国外研究现状................................................ 3',
    '   1.3 软件设计内容及目标............................................... 4',
    '2 系统分析 ................................. 5',
    '   2.1 需求分析................................................................... 5',
    '   2.2 业务流程分析........................................................... 6',
    '   2.3 功能模块分析........................................................... 7',
    '   2.4 技术分析................................................................... 10',
    '3 系统设计 ................................. 12',
    '   3.1 软件总体设计........................................................... 12',
    '   3.2 系统详细设计........................................................... 13',
    '      3.2.1 普通用户功能模块设计.................................. 13',
    '      3.2.2 管理员功能模块设计...................................... 16',
    '   3.3 数据库的设计........................................................... 18',
    '      3.3.1 概要设计.............................................................. 19',
    '      3.3.2 逻辑设计.............................................................. 24',
    '4 系统实现 ................................. 30',
    '   4.1 前端功能实现............................................................... 30',
    '      4.1.1 用户登录注册模块.............................................. 30',
    '      4.1.2 题库浏览与检索模块.......................................... 32',
    '      4.1.3 专项练习模块...................................................... 34',
    '      4.1.4 模拟面试模块...................................................... 36',
    '      4.1.5 社区圈子模块...................................................... 37',
    '   4.2 后端功能实现............................................................... 38',
    '      4.2.1 题目管理模块...................................................... 38',
    '      4.2.2 用户认证与权限管理模块.................................. 42',
    '      4.2.3 练习管理模块...................................................... 44',
    '      4.2.4 面试管理模块...................................................... 45',
    '      4.2.5 圈子管理模块...................................................... 46',
    '      4.2.6 文件存储与网关模块.......................................... 47',
    '5 系统测试 ................................. 49',
    '   5.1 系统测试概述............................................................... 49',
    '   5.2 系统测试过程............................................................... 50',
    '      5.2.1 测试用例.............................................................. 50',
    '   5.3 测试结果................................................................... 56',
    '结论 ...................................... 57',
    '参考文献 .................................. 58',
    '附录 ...................................... 60',
    '致谢 ...................................... 76',
]

for line in toc_lines:
    add_para(line, line_spacing=Pt(24), first_indent=None)
page_break()

# ================================================================
# CHAPTER 1: INTRODUCTION
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('1 绪论')

# 1.1 Background
add_section_title('1.1 软件研究背景和意义')
ch1_1 = read_content('ch1_1_background.txt')
if ch1_1:
    for para_text in ch1_1.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# 1.2 Research Status
add_section_title('1.2 软件的国内外的研究现状及发展趋势')
add_subsection_title('1.2.1 国内研究现状')
ch1_2_d = read_content('ch1_2_domestic.txt')
if ch1_2_d:
    for para_text in ch1_2_d.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

add_subsection_title('1.2.2 国外研究现状')
ch1_2_f = read_content('ch1_2_foreign.txt')
if ch1_2_f:
    for para_text in ch1_2_f.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# 1.3 Design Content
add_section_title('1.3 软件设计内容及目标')
ch1_3 = read_content('ch1_3_design_content.txt')
if ch1_3:
    for para_text in ch1_3.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

page_break()

print("Chapter 1 done...")

# ================================================================
# CHAPTER 2: SYSTEM ANALYSIS
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('2 系统分析')

# 2.1
add_section_title('2.1 需求分析')
ch2_1 = read_content('ch2_1_requirements.txt')
if ch2_1:
    for para_text in ch2_1.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# 2.2
add_section_title('2.2 业务流程分析')
ch2_2 = read_content('ch2_2_business_process.txt')
if ch2_2:
    for para_text in ch2_2.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

add_para('图2.1 系统业务流程图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

# 2.3
add_section_title('2.3 功能模块分析')
ch2_3 = read_content('ch2_3_module_analysis.txt')
if ch2_3:
    for para_text in ch2_3.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

add_para('图2.2 普通用户用例图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图2.3 管理员用例图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图2.4 系统软件结构图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

# 2.4
add_section_title('2.4 技术分析')
ch2_4 = read_content('ch2_4_tech_analysis.txt')
if ch2_4:
    for para_text in ch2_4.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

page_break()

print("Chapter 2 done...")

# ================================================================
# CHAPTER 3: SYSTEM DESIGN
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('3 系统设计')

# 3.1
add_section_title('3.1 软件总体设计')
ch3_1 = read_content('ch3_1_overall_design.txt')
if ch3_1:
    for para_text in ch3_1.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

add_para('图3.1 系统软件结构图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

# 3.2
add_section_title('3.2 系统详细设计')
ch3_2 = read_content('ch3_2_detailed_design.txt')
if ch3_2:
    for para_text in ch3_2.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

add_para('图3.2 用户登录程序序列图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.3 题目全文检索程序序列图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.4 管理员新增题目程序序列图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.5 管理员用户权限管理程序序列图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

# 3.3
add_section_title('3.3 数据库的设计')
add_subsection_title('3.3.1 概要设计')
ch3_3 = read_content('ch3_3_database.txt')
if ch3_3:
    # The first part about ER diagrams
    parts = ch3_3.split('以下为系统主要数据库表的逻辑设计')
    if len(parts) >= 1:
        for para_text in parts[0].split('\n\n'):
            para_text = para_text.strip()
            if para_text:
                add_body(para_text)

add_para('图3.6 用户管理E-R图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.7 题目管理E-R图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.8 练习管理E-R图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.9 面试管理E-R图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)
add_para('图3.10 社区圈子E-R图', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None)

# 3.3.2
add_subsection_title('3.3.2 逻辑设计')
if len(parts) >= 2:
    for para_text in parts[1].split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# Add key database tables
# auth_user table
add_para('表3.1 auth_user（用户表）', bold=True, first_indent=None)
add_table_with_data(
    ['字段名', '数据类型', '是否为空', '说明'],
    [
        ['id', 'bigint(20)', '否', '主键'],
        ['user_name', 'varchar(32)', '是', '用户名/账号'],
        ['nick_name', 'varchar(32)', '是', '昵称'],
        ['email', 'varchar(32)', '是', '邮箱'],
        ['phone', 'varchar(32)', '是', '手机号'],
        ['password', 'varchar(64)', '否', '密码（BCrypt加密）'],
        ['sex', 'tinyint(2)', '是', '性别'],
        ['avatar', 'varchar(255)', '是', '头像URL'],
        ['status', 'tinyint(2)', '是', '0启用 1禁用'],
        ['created_time', 'datetime', '是', '创建时间'],
        ['update_time', 'datetime', '是', '更新时间'],
        ['is_deleted', 'int(11)', '是', '逻辑删除标记'],
    ]
)

# subject_info table
add_para('表3.2 subject_info（题目信息表）', bold=True, first_indent=None)
add_table_with_data(
    ['字段名', '数据类型', '是否为空', '说明'],
    [
        ['id', 'bigint(20)', '否', '主键'],
        ['subject_name', 'varchar(255)', '否', '题目名称'],
        ['subject_type', 'int(11)', '否', '1单选 2多选 3判断 4简答'],
        ['subject_difficult', 'int(11)', '否', '题目难度'],
        ['subject_score', 'decimal(10,2)', '否', '题目分数'],
        ['subject_parse', 'text', '是', '题目解析'],
        ['created_by', 'varchar(32)', '是', '创建人'],
        ['created_time', 'datetime', '是', '创建时间'],
        ['is_deleted', 'int(11)', '是', '逻辑删除标记'],
    ]
)

# practice_info table
add_para('表3.3 practice_info（练习记录表）', bold=True, first_indent=None)
add_table_with_data(
    ['字段名', '数据类型', '是否为空', '说明'],
    [
        ['id', 'bigint(20)', '否', '主键'],
        ['set_id', 'bigint(20)', '是', '套题ID'],
        ['complete_status', 'int(11)', '是', '1完成 0未完成'],
        ['time_use', 'varchar(32)', '是', '用时'],
        ['submit_time', 'datetime', '是', '交卷时间'],
        ['correct_rate', 'decimal(10,2)', '是', '正确率'],
        ['created_by', 'varchar(32)', '是', '创建人'],
        ['created_time', 'datetime', '是', '创建时间'],
    ]
)

# share_circle table
add_para('表3.4 share_circle（圈子表）', bold=True, first_indent=None)
add_table_with_data(
    ['字段名', '数据类型', '是否为空', '说明'],
    [
        ['id', 'bigint(20)', '否', '主键'],
        ['circle_name', 'varchar(64)', '否', '圈子名称'],
        ['circle_desc', 'varchar(255)', '是', '圈子描述'],
        ['circle_icon', 'varchar(255)', '是', '圈子图标'],
        ['created_by', 'varchar(32)', '是', '创建者'],
        ['created_time', 'datetime', '是', '创建时间'],
        ['is_deleted', 'int(11)', '是', '逻辑删除标记'],
    ]
)

page_break()
print("Chapter 3 done...")

# ================================================================
# CHAPTER 4: SYSTEM IMPLEMENTATION
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('4 系统实现')

# 4.1 Frontend
add_section_title('4.1 前端功能实现')
ch4_1 = read_content('ch4_1_frontend.txt')
if ch4_1:
    for para_text in ch4_1.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            # Check if it's a subsection title
            if para_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
                add_subsection_title(para_text.split('\n')[0])
            add_body(para_text)

# 4.2 Backend
add_section_title('4.2 后端功能实现')
ch4_2 = read_content('ch4_2_backend.txt')
if ch4_2:
    for para_text in ch4_2.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

page_break()
print("Chapter 4 done...")

# ================================================================
# CHAPTER 5: SYSTEM TESTING
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('5 系统测试')

# 5.1
add_section_title('5.1 系统测试概述')
ch5_1 = read_content('ch5_1_testing_overview.txt')
if ch5_1:
    for para_text in ch5_1.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# 5.2
add_section_title('5.2 系统测试过程')
add_subsection_title('5.2.1 测试用例')
ch5_2 = read_content('ch5_2_test_cases.txt')
if ch5_2:
    for para_text in ch5_2.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# Test case tables
add_para('表5.1 用户注册测试用例表', bold=True, first_indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_with_data(
    ['编号', '内容', '预期结果', '测试结果', '是否相同'],
    [
        ['1', '打开注册页，输入合法用户名、密码、邮箱', '用户名正常显示，密码保密', '用户名正常显示，密码保密', '是'],
        ['2', '点击注册按钮', '提示“注册成功”', '提示“注册成功”', '是'],
        ['3', '用户名为空时点击注册', '提示“用户名不能为空”', '提示“用户名不能为空”', '是'],
        ['4', '密码长度不足6位时注册', '提示“密码长度不能少于6位”', '提示“密码长度不能少于6位”', '是'],
        ['5', '注册已存在的用户名', '提示“该用户名已被注册”', '提示“该用户名已被注册”', '是'],
    ]
)

add_para('表5.2 题目检索测试用例表', bold=True, first_indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_with_data(
    ['编号', '内容', '预期结果', '测试结果', '是否相同'],
    [
        ['1', '选择“Java”分类和“多线程”标签', '展示Java分类下多线程标签的题目列表', '正确展示符合条件的题目列表', '是'],
        ['2', '搜索“线程池”关键词', '返回包含“线程池”的题目，高亮关键词', '正确返回搜索结果，高亮匹配关键词', '是'],
        ['3', '搜索不存在的关键词', '提示“未找到相关题目”', '提示“未找到相关题目”，未报错', '是'],
    ]
)

add_para('表5.3 智能组卷测试用例表', bold=True, first_indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_with_data(
    ['编号', '内容', '预期结果', '测试结果', '是否相同'],
    [
        ['1', '设定10题，难度分布均匀，选择“数据结构”分类', '生成符合要求的试卷', '正确生成试卷，题目数量和难度分布符合设定', '是'],
        ['2', '提交答案，查看评分和解析', '显示正确率、每题解析', '正确评分，解析准确', '是'],
        ['3', '分类下题目不足时设定超过实际数量', '提示“题目数量不足”', '提示“该分类下题目数量不足，请调整参数”', '是'],
        ['4', '设定题目总数为0', '提示“题目数量不能为0”', '提示“题目数量不能为0”', '是'],
    ]
)

# 5.3
add_section_title('5.3 测试结果')
ch5_3 = read_content('ch5_3_test_results.txt')
if ch5_3:
    for para_text in ch5_3.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

page_break()
print("Chapter 5 done...")

# ================================================================
# CONCLUSION
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('结论')

conclusion = read_content('conclusion.txt')
if conclusion:
    for para_text in conclusion.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

page_break()

# ================================================================
# REFERENCES
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('参考文献')

refs = read_content('references.txt')
if refs:
    for ref_line in refs.split('\n'):
        ref_line = ref_line.strip()
        if ref_line:
            add_para(ref_line, first_indent=Cm(0), font_west='Times New Roman',
                    line_spacing=Pt(22), size=Pt(10.5))

page_break()

# ================================================================
# APPENDIX
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('附录')

appendix_files = [
    ('appendix_1_login.txt', '附录一：'),
    ('appendix_2_subject.txt', '附录二：'),
    ('appendix_3_practice.txt', '附录三：'),
    ('appendix_4_auth.txt', '附录四：'),
    ('appendix_5_circle.txt', '附录五：'),
    ('appendix_6_gateway.txt', '附录六：'),
]

for filename, label in appendix_files:
    content = read_content(filename)
    if content:
        # Split content into text and code blocks
        # First line is the appendix title/label
        lines = content.split('\n')
        title_line = lines[0].strip()
        if title_line:
            add_body(title_line)

        # Process remaining content: extract text paragraphs and code blocks
        remaining = '\n'.join(lines[1:])

        # Split by code block markers ```
        parts = remaining.split('```')

        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue
            if i % 2 == 1:
                # Odd parts are code (between ``` markers)
                # Remove language identifier line if present (e.g. "java", "yaml")
                code_lines = part.split('\n')
                if code_lines and code_lines[0].strip() in ('java', 'python', 'typescript', 'yaml', 'xml', 'json'):
                    code_lines = code_lines[1:]
                code_text = '\n'.join(code_lines).strip()
                if code_text:
                    add_code(code_text)
            else:
                # Even parts are text between code blocks
                for para_text in part.split('\n\n'):
                    para_text = para_text.strip()
                    if para_text and len(para_text) > 10:
                        add_body(para_text)
        add_empty_line()

page_break()
print("Appendix done...")

# ================================================================
# ACKNOWLEDGEMENT
# ================================================================
add_para('江西农业大学南昌商学院',
         font_east='宋体', size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
add_chapter_title('致谢')

ack = read_content('acknowledgement.txt')
if ack:
    for para_text in ack.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            add_body(para_text)

# ================================================================
# SAVE
# ================================================================
output_path = os.path.join(os.path.dirname(__file__), 'jc-club-thesis.docx')
doc.save(output_path)
print(f"Thesis saved to: {output_path}")
print("Done! Complete thesis generated successfully.")
